import os
import re
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

def clean_cell_text(text):
    if text is None:
        return ""
    # Replace non-breaking spaces and clean whitespace
    text = text.replace('\xa0', ' ').strip()
    return text

def convert_val(text):
    cleaned = clean_cell_text(text)
    if not cleaned:
        return None
    # Try converting to int or float
    try:
        if re.match(r'^-?\d+$', cleaned):
            return int(cleaned)
        elif re.match(r'^-?\d*\.\d+$', cleaned):
            return float(cleaned)
    except ValueError:
        pass
    return cleaned

def convert_docx_to_xlsx(docx_path, xlsx_path):
    print(f"Loading document: {docx_path}")
    doc = Document(docx_path)
    
    if not doc.tables:
        raise ValueError("No tables found in the document!")
    
    table = doc.tables[0]
    print(f"Found table with {len(table.rows)} rows and {len(table.columns)} columns.")
    
    wb = Workbook()
    
    # Set publication metadata properties
    wb.properties.title = "Oases of Endemism: Data Table S5"
    wb.properties.creator = "Matthew J. Forrest"
    wb.properties.description = (
        "Dataset corresponding to Data Table S5 from the article:\n"
        "Oases of endemism: Regional aquifer desert springs serve as biodiversity hotspots "
        "preserving vulnerable endemic taxa in the Great Basin and Mojave Desert regions.\n"
        "Published in Limnology and Oceanography (2026). DOI: 10.1002/lno.70414"
    )
    wb.properties.subject = "Limnology, Hydrology, Desert Springs, Biodiversity, Endemic Taxa"
    wb.properties.category = "Scientific Dataset"
    
    ws = wb.active
    ws.title = "Data Table S5"
    
    # Ensure grid lines are visible in Excel
    ws.views.sheetView[0].showGridLines = True
    
    # Styling definitions
    header_font = Font(name="Calibri", size=11, bold=True)
    body_font = Font(name="Calibri", size=11)
    
    header_fill = PatternFill(start_color="EAEAEA", end_color="EAEAEA", fill_type="solid")
    
    thin_border = Border(
        left=Side(style='thin', color='D3D3D3'),
        right=Side(style='thin', color='D3D3D3'),
        top=Side(style='thin', color='D3D3D3'),
        bottom=Side(style='thin', color='D3D3D3')
    )
    
    print("Extracting and writing rows to Excel...")
    for r_idx, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            # Extract text from cell paragraphs
            cell_text = "\n".join([p.text for p in cell.paragraphs]).strip()
            row_data.append(cell_text)
            
        # Convert values based on row type
        converted_row = []
        for val in row_data:
            if r_idx < 2:
                # Row 0 (Variable names) and Row 1 (Units) - keep as text
                converted_row.append(clean_cell_text(val))
            else:
                converted_row.append(convert_val(val))
                
        ws.append(converted_row)
        
        # Style the written row
        excel_row_idx = r_idx + 1
        for c_idx in range(1, len(converted_row) + 1):
            cell = ws.cell(row=excel_row_idx, column=c_idx)
            cell.font = header_font if r_idx < 2 else body_font
            if r_idx < 2:
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border = thin_border
            else:
                cell.border = thin_border
                val = cell.value
                if isinstance(val, (int, float)):
                    cell.alignment = Alignment(horizontal="right")
                else:
                    cell.alignment = Alignment(horizontal="left")
                    
    # Adjust column widths dynamically
    print("Formatting column widths...")
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            val_str = str(cell.value or '')
            # Split by newline for multi-line cells to get actual visual width
            lines = val_str.split('\n')
            for line in lines:
                if len(line) > max_len:
                    max_len = len(line)
        # Apply padding and set column width
        ws.column_dimensions[col_letter].width = max(max_len + 3, 10)
        
    print(f"Saving workbook to: {xlsx_path}")
    wb.save(xlsx_path)
    print("Conversion complete successfully!")

if __name__ == "__main__":
    convert_docx_to_xlsx("Data_Table_S5.docx", "Data_Table_S5.xlsx")
