import docx

def main():
    print("Reading Data_Table_S5.docx...")
    doc = docx.Document("Data_Table_S5.docx")
    
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    print(f"Number of tables: {len(doc.tables)}")
    
    # Search paragraphs
    for idx, p in enumerate(doc.paragraphs):
        text = p.text
        for term in ["AIC", "BIC", "model", "best-fit", "SIMPER", "simprer", "S1", "S2"]:
            if term.lower() in text.lower():
                print(f"Found in paragraph {idx}: ... {text[:150]} ...")
                break
                
    # Check tables (beyond the first large data table)
    # The first table is the main dataset, let's check if there are other tables
    if len(doc.tables) > 1:
        for t_idx, t in enumerate(doc.tables[1:], 1):
            print(f"Checking table {t_idx}...")
            # print first row to see header
            row_vals = [cell.text.strip() for cell in t.rows[0].cells]
            print(f"  Header: {row_vals}")

if __name__ == "__main__":
    main()
